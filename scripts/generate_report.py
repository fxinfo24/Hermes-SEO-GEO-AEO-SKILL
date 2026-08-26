#!/usr/bin/env python3
"""
SEO / GEO / AEO Audit report generator (Hermes port).

Replaces the original Claude-only `node docx` script. Accepts a JSON audit
file (recommended) or falls back to the built-in SAMPLE_DATA for a demo.

Usage:
    python3 generate_report.py                 # demo using built-in SAMPLE_DATA
    python3 generate_report.py audit.json      # render from JSON
    python3 generate_report.py audit.json --out my-report.docx

Design system is ported verbatim from SKILL.md:
  Navy 1B2A4A, Accent 2563EB, green 16A34A, amber D97706, red DC2626,
  gray rows F8F9FA, borders E2E8F0, dark text 1E293B, light bg EFF6FF.

JSON schema (all keys except `glossary` / `health_score` required):
{
  "domain": "example.com",
  "audit_type": "FULL AUDIT",                 # "QUICK AUDIT" or "FULL AUDIT"
  "date": "2025-03-13",
  "scores": {"SEO": 7, "GEO": 6, "AEO": 5},   # each 1-10
  "takeaways": {"SEO": "...", "GEO": "...", "AEO": "...", "Combined": "..."},
  "health_score": 73,                         # OPTIONAL 0-100 (shows on cover + Combined row)
  "author": "The Saint",                      # OPTIONAL; shown on cover (defaults to skill author)
  "exec_summary": "...",
  "pages_audited": [["url", "type", "notes"], ...],
  "seo":  [["Subsection", [["signal","finding","status"], ...]], ...],
  "geo":  [["Subsection", [["signal","finding","status"], ...]], ...],
  "aeo":  [["Subsection", [["signal","finding","status"], ...]], ...],
  "recommendations": [["Critical","issue","dim","Low","High"], ...],
  "strengths": [["strength","evidence"], ...],
  "glossary": true                            # Full Audit -> include glossary
}
status in {Good, Needs Attention, Missing}; priority in {Critical, High, Medium, Quick Win}
"""

import argparse
import json
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Design tokens
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x25, 0x63, 0xEB)
LIGHT_BLUE = RGBColor(0x93, 0xC5, 0xFD)
GRAY = RGBColor(0x94, 0xA3, 0xB8)
DARK = RGBColor(0x1E, 0x29, 0x3B)

C_NAVY = "1B2A4A"
C_ACCENT = "2563EB"
C_GREEN = "16A34A"
C_AMBER = "D97706"          # amber == High priority (matches SKILL.md)
C_RED = "DC2626"
C_ROW = "F8F9FA"
C_BORDER = "E2E8F0"
C_LIGHTBG = "EFF6FF"
C_GREENBG = "F0FDF4"

FONT = "Arial"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def score_color(n, scale=10):
    if scale == 100:
        return C_GREEN if n >= 80 else (C_AMBER if n >= 50 else C_RED)
    if n >= 8:
        return C_GREEN
    if n >= 5:
        return C_AMBER
    return C_RED


def score_status(n):
    if n >= 8:
        return "Strong"
    if n >= 5:
        return "On Track"
    return "Needs Work"


def status_color(s):
    return {"Good": C_GREEN, "Needs Attention": C_AMBER, "Missing": C_RED}.get(s, C_AMBER)


def priority_color(p):
    # High is amber D97706 (per SKILL.md); Critical red, Medium amber, Quick Win green
    return {"Critical": C_RED, "High": C_AMBER, "Medium": C_AMBER,
            "Quick Win": C_GREEN}.get(p, C_AMBER)


def style_run(r, size, bold=False, color=None, italic=False):
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color


def add_heading(doc, text, level):
    sizes = {1: 24, 2: 18, 3: 14}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text)
    style_run(r, sizes.get(level, 14), bold=True, color=NAVY)
    return p


def filled_cell(cell, text, size=10, bold=False, color=None, align="left", fill=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(text)
    style_run(r, size, bold=bold, color=color, italic=italic)
    set_cell_margins(cell)
    if fill:
        shade(cell, fill)
    return cell


def add_row_to(table, ncols):
    """python-docx only materializes cell 0 of row 0; add_row() for the rest."""
    row = table.add_row()
    return [row.cells[c] for c in range(ncols)]


def add_code_block(doc, code, label=None):
    """Render a monospace, shaded code block (for copy-paste snippets)."""
    if label:
        lp = doc.add_paragraph()
        lr = lp.add_run(label)
        style_run(lr, 10, bold=True, color=ACCENT)
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    c = box.rows[0].cells[0]
    shade(c, "F4F6F8")
    set_cell_margins(c, top=120, bottom=120, left=160, right=160)
    c.text = ""
    first = True
    for line in code.strip("\n").split("\n"):
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = DARK
    doc.add_paragraph()


def add_fixit_section(doc, fixit):
    """Render the 'How to Fix' appendix from DATA['fixit'] (list of
    {title, label?, code} dicts)."""
    add_heading(doc, "How to Fix (copy-paste)", 1)
    for item in fixit:
        add_heading(doc, item.get("title", "Snippet"), 2)
        add_code_block(doc, item["code"], label=item.get("label"))


def add_aeo_section(doc, aeo):
    """Render an AEO before/after example block."""
    RED = RGBColor(0xDC, 0x26, 0x26)
    GREEN = RGBColor(0x16, 0xA3, 0x4A)
    add_heading(doc, "AEO Before / After Example", 1)
    for ex in aeo:
        add_heading(doc, ex.get("page", "Example"), 2)
        p = doc.add_paragraph()
        r = p.add_run("BEFORE — declarative heading, no direct answer:")
        style_run(r, 10, bold=True, color=RED)
        add_code_block(doc, ex.get("before", ""))
        p = doc.add_paragraph()
        r = p.add_run("AFTER — question heading + 40-60 word direct answer + FAQ schema:")
        style_run(r, 10, bold=True, color=GREEN)
        add_code_block(doc, ex.get("after", ""))


# ---------------------------------------------------------------------------
# Validation + loading
# ---------------------------------------------------------------------------
def validate(data):
    required = ["domain", "audit_type", "date", "scores", "exec_summary",
                "pages_audited", "seo", "geo", "aeo", "recommendations", "strengths"]
    missing = [k for k in required if k not in data]
    if missing:
        raise ValueError(f"audit JSON missing required keys: {', '.join(missing)}")
    for d in ("SEO", "GEO", "AEO"):
        if d not in data["scores"]:
            raise ValueError(f"scores must include '{d}'")
    for key in ("seo", "geo", "aeo", "pages_audited", "recommendations", "strengths"):
        if not isinstance(data[key], list):
            raise ValueError(f"'{key}' must be a list")
    # optional keys (no schema enforcement beyond presence):
    #   takeaways, health_score, glossary, fixit, aeo_example
    return data


def load_audit(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except FileNotFoundError:
        raise SystemExit(f"ERROR: audit file not found: {path}")
    except json.JSONDecodeError as e:
        raise SystemExit(f"ERROR: invalid JSON in {path}: {e}")
    return validate(raw)


# ---------------------------------------------------------------------------
# Built-in demo data (used when no JSON file is supplied)
# ---------------------------------------------------------------------------
SAMPLE_DATA = {
    "domain": "example.com",
    "audit_type": "FULL AUDIT",
    "date": "2025-03-13",
    "author": "The Saint",
    "scores": {"SEO": 7, "GEO": 6, "AEO": 5},
    "takeaways": {
        "SEO": "Technical foundation solid; meta descriptions missing on key pages.",
        "GEO": "Strong E-E-A-T; thin structured-data depth (no Author/Dataset).",
        "AEO": "No FAQ/HowTo schema; few question-phrased headings.",
        "Combined": "Good base with large AI-search upside.",
    },
    "health_score": 73,
    "exec_summary": (
        "Example.com has a solid technical foundation and strong E-E-A-T signals "
        "on its About page, but is under-optimized for AI search (GEO) and misses "
        "featured-snippet opportunities (AEO). The most urgent issue is the absence "
        "of FAQ and HowTo schema. The key opportunity is restructuring service "
        "pages around question-phrased headings to capture AI Overviews."
    ),
    "pages_audited": [
        ["https://example.com/", "Homepage", "Clear H1; missing meta description"],
        ["https://example.com/about/", "About / Team", "Strong author credentials"],
        ["https://example.com/services/", "Services", "Rich content; no FAQ schema"],
        ["https://example.com/contact/", "Contact", "NAP present; no LocalBusiness schema"],
    ],
    "seo": [
        ("Technical On-Page", [
            ("Title tag", "Present, 52 chars, contains primary keyword.", "Good"),
            ("Meta description", "Missing on homepage.", "Missing"),
            ("Heading hierarchy", "Single H1; logical H2/H3.", "Good"),
            ("Canonical tag", "Self-referencing correctly.", "Good"),
            ("Image alt text", "Several decorative images lack alt text.", "Needs Attention"),
        ]),
        ("Content Quality", [
            ("Word count", "Service pages ~900 words; pillar content 1,600+.", "Good"),
            ("Keyword signals", "Primary topic clear; thin semantic coverage.", "Needs Attention"),
            ("Freshness", "No visible publish/update dates.", "Needs Attention"),
        ]),
        ("Structured Data", [
            ("Schema markup", "Organization JSON-LD present; no FAQ/HowTo.", "Needs Attention"),
        ]),
    ],
    "geo": [
        ("E-E-A-T Assessment", [
            ("Author information", "Named authors with credentials on About.", "Good"),
            ("Organization schema", "Declared with logo + URL.", "Good"),
            ("Trust signals", "Testimonials present; no awards/press.", "Needs Attention"),
        ]),
        ("Content for AI Synthesis", [
            ("Factual density", "Specific stats cited in case studies.", "Good"),
            ("Source citation", "Few external authoritative citations.", "Needs Attention"),
        ]),
        ("Technical GEO", [
            ("Structured data depth", "No Author/Dataset/ClaimReview types.", "Needs Attention"),
            ("HTTPS", "Site is fully HTTPS.", "Good"),
        ]),
    ],
    "aeo": [
        ("Featured Snippet Eligibility", [
            ("Direct answer paragraphs", "No 40-60 word answer blocks under Q headings.", "Missing"),
            ("Definition patterns", "Core topics not stated as 'X is...' sentences.", "Needs Attention"),
        ]),
        ("Structured Answer Formats", [
            ("FAQ schema", "No FAQ schema anywhere.", "Missing"),
            ("HowTo schema", "Step content not marked up with HowTo.", "Missing"),
            ("Question-phrased headings", "Headings are declarative, not questions.", "Needs Attention"),
        ]),
        ("Voice Search Readiness", [
            ("Conversational language", "Mostly formal; some long-tail coverage.", "Needs Attention"),
            ("Local signals", "NAP on contact; no local schema.", "Needs Attention"),
        ]),
    ],
    "recommendations": [
        ("Critical", "Add FAQ schema to Services and Blog pages.", "AEO", "Low", "High"),
        ("High", "Write unique meta descriptions for all top pages.", "SEO", "Low", "Medium"),
        ("Medium", "Restructure service pages with question-phrased H2s.", "GEO/AEO", "Medium", "High"),
        ("Quick Win", "Add alt text to decorative/hero images.", "SEO", "Low", "Low"),
    ],
    "strengths": [
        ("Strong About/Team page", "Named authors with real credentials build E-E-A-T."),
        ("Secure, fast-foundation site", "Full HTTPS and clean URL structure."),
        ("Pillar content depth", "Service guides exceed 1,500 words with specifics."),
    ],
    "glossary": True,
}


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build_document(DATA):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = DARK

    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # ---- Cover page (navy canvas via 1x1 table) ----
    ct = doc.add_table(rows=1, cols=1)
    ct.allow_autofit = False
    ct.columns[0].width = Inches(6.5)
    cell = ct.rows[0].cells[0]
    shade(cell, C_NAVY)
    set_cell_margins(cell, top=1800, bottom=1800, left=400, right=400)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DATA["content"]["domain"] if "content" in DATA else DATA["domain"])
    style_run(r, 36, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SEO / GEO / AEO Audit Report")
    style_run(r2, 18, color=LIGHT_BLUE)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(DATA["audit_type"])
    style_run(r3, 11, color=RGBColor(0xFF, 0xFF, 0xFF))
    p3.paragraph_format.space_after = Pt(20)

    st_tbl = cell.add_table(rows=3, cols=3)
    st_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_borders(st_tbl)
    for ci, d in enumerate(["SEO", "GEO", "AEO"]):
        n = DATA["scores"][d]
        fill = score_color(n)
        for c in (st_tbl.cell(0, ci), st_tbl.cell(1, ci), st_tbl.cell(2, ci)):
            shade(c, fill)
            set_cell_margins(c, top=120, bottom=120)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c0 = st_tbl.cell(0, ci); c0.text = ""
        rr = c0.paragraphs[0].add_run(d)
        style_run(rr, 10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c1 = st_tbl.cell(1, ci); c1.text = ""
        rr = c1.paragraphs[0].add_run(f"{n}")
        style_run(rr, 36, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        c2 = st_tbl.cell(2, ci); c2.text = ""
        rr = c2.paragraphs[0].add_run(score_status(n))
        style_run(rr, 9, italic=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # optional overall health score (0-100) on cover
    if DATA.get("health_score") is not None:
        ph = cell.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ph.paragraph_format.space_before = Pt(14)
        rh = ph.add_run(f"Overall Health Score: {DATA['health_score']}/100")
        style_run(rh, 12, bold=True, color=LIGHT_BLUE)

    pa = cell.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa.paragraph_format.space_before = Pt(20)
    author = DATA.get("author") or "The Saint"
    ra = pa.add_run(f"Audit date: {DATA['date']}\nHermes SEO/GEO/AEO Skill\nAuthor: {author}")
    style_run(ra, 9, color=GRAY)
    doc.add_page_break()

    # ---- Header / footer ----
    hp = sec.header.paragraphs[0]
    hp.text = ""
    hr = hp.add_run(f"{DATA['domain']}")
    style_run(hr, 9, color=DARK)
    htab = hp.add_run("\tSEO / GEO / AEO Audit Report")
    htab.font.name = FONT
    htab.font.size = Pt(9)
    hPr = sec.header.paragraphs[0]._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), C_NAVY)
    pbdr.append(bottom)
    hPr.append(pbdr)

    fp = sec.footer.paragraphs[0]
    fp.text = ""
    fr = fp.add_run("Hermes SEO/GEO/AEO Skill")
    style_run(fr, 9, color=GRAY)
    ftp = fp.add_run("\t\t")
    fn = fp.add_run()
    fld1 = OxmlElement("w:fldSimple")
    fld1.set(qn("w:instr"), "PAGE")
    fn._r.append(fld1)
    fn.font.name = FONT
    fn.font.size = Pt(9)
    fn.font.color.rgb = GRAY
    fPr = sec.footer.paragraphs[0]._p.get_or_add_pPr()
    fbdr = OxmlElement("w:pBdr")
    ftop = OxmlElement("w:top")
    ftop.set(qn("w:val"), "single")
    ftop.set(qn("w:sz"), "6")
    ftop.set(qn("w:space"), "1")
    ftop.set(qn("w:color"), C_BORDER)
    fbdr.append(ftop)
    fPr.append(fbdr)

    # ---- 2. Executive Summary ----
    add_heading(doc, "Executive Summary", 1)
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    bc = box.rows[0].cells[0]
    shade(bc, C_LIGHTBG)
    set_cell_margins(bc, top=160, bottom=160, left=200, right=200)
    bc.text = ""
    bp = bc.paragraphs[0]
    br = bp.add_run(DATA["exec_summary"])
    style_run(br, 11, color=DARK)
    bp.paragraph_format.space_after = Pt(6)

    add_heading(doc, "Scores", 3)

    seo_n, geo_n, aeo_n = DATA["scores"]["SEO"], DATA["scores"]["GEO"], DATA["scores"]["AEO"]
    takeaways = DATA.get("takeaways", {})
    if DATA.get("health_score") is not None:
        combined_n = DATA["health_score"]
        combined_scale = 100
        combined_disp = f"{combined_n}/100"
        combined_status = ""
    else:
        combined_n = seo_n + geo_n + aeo_n
        combined_scale = 10
        combined_disp = f"{combined_n}/30"
        combined_status = score_status(combined_n)

    rows = [
        ("SEO", seo_n, f"{seo_n}/10", score_status(seo_n), 10, takeaways.get("SEO", "")),
        ("GEO", geo_n, f"{geo_n}/10", score_status(geo_n), 10, takeaways.get("GEO", "")),
        ("AEO", aeo_n, f"{aeo_n}/10", score_status(aeo_n), 10, takeaways.get("AEO", "")),
        ("Combined", combined_n, combined_disp, combined_status, combined_scale, takeaways.get("Combined", "")),
    ]
    sc = doc.add_table(rows=1, cols=4)
    sc.style = "Table Grid"
    for ci, h in enumerate(["Dimension", "Score", "Status", "Key Takeaway"]):
        filled_cell(sc.rows[0].cells[ci], h, size=11, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), fill=C_NAVY)
    for (d, n, s, st_, scale, kt) in rows:
        fill = score_color(n, scale) if n is not None else None
        rc = add_row_to(sc, 4)
        filled_cell(rc[0], d, size=10, bold=True)
        if fill:
            filled_cell(rc[1], s, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                        align="center", fill=fill)
        else:
            filled_cell(rc[1], s, size=10, bold=True, align="center")
        filled_cell(rc[2], st_, size=10, align="center")
        filled_cell(rc[3], kt, size=10)

    # ---- 3. Pages Audited ----
    add_heading(doc, "Pages Audited", 1)
    pat = doc.add_table(rows=1, cols=3)
    pat.style = "Table Grid"
    for ci, h in enumerate(["URL", "Page Type", "Notes"]):
        filled_cell(pat.rows[0].cells[ci], h, size=11, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), fill=C_NAVY)
    for i, row in enumerate(DATA["pages_audited"], start=1):
        url, ptype, notes = row[0], row[1], row[2]
        fill = C_ROW if i % 2 == 0 else None
        rc = add_row_to(pat, 3)
        filled_cell(rc[0], url, size=9, fill=fill)
        filled_cell(rc[1], ptype, size=9, fill=fill)
        filled_cell(rc[2], notes, size=9, fill=fill)

    # ---- 4/5/6. Analysis sections ----
    for title, key in [("SEO Analysis", "seo"), ("GEO Analysis", "geo"), ("AEO Analysis", "aeo")]:
        add_heading(doc, title, 1)
        for sub, findings in DATA[key]:
            add_heading(doc, sub, 2)
            t = doc.add_table(rows=1, cols=3)
            t.style = "Table Grid"
            for ci, h in enumerate(["Signal", "Finding", "Status"]):
                filled_cell(t.rows[0].cells[ci], h, size=10, bold=True,
                            color=RGBColor(0xFF, 0xFF, 0xFF), fill=C_NAVY)
            for i, frow in enumerate(findings, start=1):
                sig, find, stt = frow[0], frow[1], frow[2]
                fill = C_ROW if i % 2 == 0 else None
                rc = add_row_to(t, 3)
                filled_cell(rc[0], sig, size=9, bold=True, fill=fill)
                filled_cell(rc[1], find, size=9, fill=fill)
                filled_cell(rc[2], stt, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                            align="center", fill=status_color(stt))

    # ---- 7. Priority Recommendations ----
    add_heading(doc, "Priority Recommendations", 1)
    rt = doc.add_table(rows=1, cols=5)
    rt.style = "Table Grid"
    for ci, h in enumerate(["Priority", "Issue", "Dimension", "Effort", "Impact"]):
        filled_cell(rt.rows[0].cells[ci], h, size=10, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), fill=C_NAVY)
    for row in DATA["recommendations"]:
        pri, issue, dim, eff, imp = row[0], row[1], row[2], row[3], row[4]
        rc = add_row_to(rt, 5)
        filled_cell(rc[0], pri, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
                    align="center", fill=priority_color(pri))
        filled_cell(rc[1], issue, size=9)
        filled_cell(rc[2], dim, size=9, align="center")
        filled_cell(rc[3], eff, size=9, align="center")
        filled_cell(rc[4], imp, size=9, align="center")

    # ---- 8. What's Working Well ----
    add_heading(doc, "What's Working Well", 1)
    wt = doc.add_table(rows=1, cols=2)
    wt.style = "Table Grid"
    for ci, h in enumerate(["Strength", "Evidence"]):
        filled_cell(wt.rows[0].cells[ci], h, size=10, bold=True,
                    color=RGBColor(0xFF, 0xFF, 0xFF), fill=C_GREEN)
    for i, row in enumerate(DATA["strengths"], start=1):
        s, ev = row[0], row[1]
        fill = C_GREENBG if i % 2 == 1 else None
        rc = add_row_to(wt, 2)
        filled_cell(rc[0], s, size=9, bold=True, fill=fill)
        filled_cell(rc[1], ev, size=9, fill=fill)

    # ---- 9. Glossary (Full Audit) ----
    if DATA.get("glossary"):
        add_heading(doc, "Glossary", 1)
        for term, desc in [
            ("SEO", "Search Engine Optimization — improving visibility in traditional search (Google, Bing)."),
            ("GEO", "Generative Engine Optimization — optimizing for AI search engines (Perplexity, ChatGPT Search, Google AI Overviews, Gemini)."),
            ("AEO", "Answer Engine Optimization — optimizing for featured snippets, People Also Ask, and voice search."),
        ]:
            p = doc.add_paragraph()
            r = p.add_run(f"{term}: ")
            style_run(r, 11, bold=True, color=NAVY)
            r2 = p.add_run(desc)
            style_run(r2, 11, color=DARK)

    # ---- 10. How to Fix (copy-paste) + AEO before/after ----
    if DATA.get("fixit"):
        add_fixit_section(doc, DATA["fixit"])
    if DATA.get("aeo_example"):
        add_aeo_section(doc, DATA["aeo_example"])

    domain_hyphen = DATA["domain"].replace(".", "-")
    out = f"seo-audit-{domain_hyphen}-{DATA['date']}.docx"
    return doc, out


def main():
    parser = argparse.ArgumentParser(description="SEO/GEO/AEO audit report generator")
    parser.add_argument("json_file", nargs="?", help="audit JSON file (omit for built-in demo)")
    parser.add_argument("--out", help="output .docx path (default: seo-audit-<domain>-<date>.docx)")
    args = parser.parse_args()

    DATA = load_audit(args.json_file) if args.json_file else SAMPLE_DATA

    doc, default_out = build_document(DATA)
    out = args.out or default_out
    doc.save(out)
    print(f"DOCX written: {out}")


if __name__ == "__main__":
    main()
